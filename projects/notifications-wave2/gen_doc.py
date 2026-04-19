from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# === LANDSCAPE ORIENTATION ===
section = doc.sections[0]
section.page_width, section.page_height = section.page_height, section.page_width
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)


def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_header_row(table, texts, bg='2E74B5'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = text
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)


def add_data_row(table, texts, bg=None):
    row = table.add_row()
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = text
        if bg:
            shade_cell(cell, bg)
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].font.size = Pt(8.5)
    return row


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(9.5)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run('\u26a0 Примечание: ')
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
    r2 = p.add_run(text)
    r2.font.size = Pt(8.5)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return p


# ===================== TITLE =====================
title = doc.add_heading('Система email-уведомлений по обращениям', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph(
    'Описание работающего механизма (CTI / BPMSoft 1.9). '
    'Источники: анализ кода + документация BPMSoft v1.9 + выгрузка vwprocesslib с прода (15.04.2026)'
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
sub.runs[0].font.size = Pt(10)
d = doc.add_paragraph('Дата: 15.04.2026 (v2)  |  Архив прода: CTI от 2026-04-11')
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.runs[0].font.size = Pt(8.5)
d.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
doc.add_paragraph()

# ===================== РАЗДЕЛ 0 =====================
add_h(doc, '0. Кому отправляются уведомления — логика получателей', 1)
add_body(doc,
    'В обращении есть два разных контакта на стороне клиента:\n'
    '  \u2022 \u00abИнициатор\u00bb — контакт, который зарегистрировал обращение (кто написал/позвонил)\n'
    '  \u2022 \u00abПользователь услуги\u00bb — контакт, для которого зарегистрировано обращение (кто фактически пострадал)\n\n'
    'Это могут быть разные люди. Например: секретарь (инициатор) подала обращение за директора (пользователь услуги).\n\n'
    'Контактное лицо (кому реально уйдёт письмо) определяется системными настройками:'
)
doc.add_paragraph()

t0 = doc.add_table(rows=1, cols=4)
t0.style = 'Table Grid'
add_header_row(t0, ['Системная настройка', 'Код', 'По умолчанию', 'Что означает'])
sysset_rows = [
    ('Оповещение пользователя услуги как контактного лица',
     'DefaultNotifyContactInCase', 'ВКЛЮЧЕНА',
     'Письма клиенту уходят на \u00abПользователь услуги\u00bb.\nПо умолчанию это главный получатель.'),
    ('Оповещение инициатора как контактного лица',
     'DefaultNotifyInitiatorInCase', 'ВЫКЛЮЧЕНА',
     'Письма клиенту уходят на \u00abИнициатора\u00bb.\nПо умолчанию выключено — инициатор не получает.'),
]
for i, r in enumerate(sysset_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t0, r, bg)
set_col_widths(t0, [Cm(6), Cm(5.5), Cm(3), Cm(11)])

doc.add_paragraph()
add_note(doc,
    'Если обе настройки включены — письма уходят обоим. Если инициатор = пользователь услуги — '
    'кнопки выбора контактного лица на странице обращения не отображаются. '
    'На странице обращения инженер может вручную переключить контактное лицо (кнопки рядом с полями).\n'
    'Одна из настроек должна быть включена — иначе уведомления не будут отправляться никому.'
)
add_note(doc,
    'При регистрации обращения из входящего email: уведомление о регистрации получают ВСЕ адресаты '
    'исходного письма (кому/копия). Поле \u00abОт кого\u00bb = адрес почтового ящика службы поддержки, '
    'на который пришло письмо. Если обращение создано иначе (звонок, вручную) — \u00abОт кого\u00bb '
    'берётся из системной настройки \u00abE-mail службы поддержки\u00bb.'
)

# ===================== РАЗДЕЛ 1 =====================
doc.add_page_break()
add_h(doc, '1. Уведомления в сторону Заказчика (клиента)', 1)
add_body(doc,
    'Все уведомления клиенту при смене статуса — через C#-класс SendEmailToCaseStatusChanged '
    '(feature-toggle SendEmailToCaseOnStatusChangeClass=1), который читает справочник CaseNotificationRule '
    'и подбирает шаблон по паре (статус + категория). '
    'Для каждого события есть ДВА шаблона: один для Пользователя услуги, другой для Инициатора. '
    'Исключение — регистрация: отдельный BPMN-процесс UsrProcess_send_reg_mail.'
)
doc.add_paragraph()

t1 = doc.add_table(rows=1, cols=6)
t1.style = 'Table Grid'
add_header_row(t1, ['Событие', 'Получатель', 'Шаблон (Пользователь услуги)', 'Шаблон (Инициатор)', 'Механизм (системное имя)', 'Тип механизма'])

client_rows = [
    ('Регистрация обращения',
     'Контактное лицо\n(по настройке)',
     'Создание нового обращения из чата',
     '\u2014 (отдельного нет)',
     'UsrProcess_send_reg_mail',
     'BPMN-процесс\n(CTI)'),
    ('\u2192 «В работе»',
     'Контактное лицо\n(по настройке)',
     'Сообщение о взятии обращения в работу',
     'Сообщение о взятии обращения в работу (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('Ответ инженера по обращению',
     'Контактное лицо\n(по настройке)',
     'Сообщение о получении нового ответа по обращению',
     'Сообщение о получении нового ответа по обращению (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('\u2192 «Решено»',
     'Контактное лицо\n(по настройке)',
     'Сообщение о разрешении обращения\nИЛИ\nСообщение о разрешении — только с решением',
     'Сообщение о разрешении обращения (для инициатора)\nИЛИ\nСообщение о разрешении — только с решением (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('\u2192 «Закрыто»',
     'Контактное лицо\n(по настройке)',
     'Сообщение о закрытии обращения',
     'Сообщение о закрытии обращения (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('Закрытие по таймеру\n(долго нет ответа)',
     'Контактное лицо\n(по настройке)',
     'Подтверждение закрытия обращения',
     'Подтверждение закрытия обращения (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('\u2192 «Отменено»',
     'Контактное лицо\n(по настройке)',
     'Сообщение об отмене обращения',
     'Сообщение об отмене обращения (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('Запрос оценки\n(после решения)',
     'Контактное лицо\n(по настройке)',
     'Запрос оценки по обращению',
     'Запрос оценки по обращению (для инициатора)',
     'SendEmailToCaseStatusChanged\n+ CaseNotificationRule',
     'C#-класс (CaseService)\n+ справочник'),
    ('Сообщение с портала',
     'Контактное лицо\n(по настройке)',
     '(по шаблону)',
     '\u2014',
     'CasePortalMessageHistoryNotification\nProcessMultiLanguage',
     'BPMN (Portal)'),
]

for i, r in enumerate(client_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t1, r, bg)

set_col_widths(t1, [Cm(3.2), Cm(2.8), Cm(6.5), Cm(6.5), Cm(5.5), Cm(3.5)])
doc.add_paragraph()
add_note(doc,
    'CC: к письмам клиенту автоматически добавляются CC-адреса из контракта и аккаунта — '
    'кастомизация через UsrActivityCcEventListener (CTI).\n'
    'Поле «Процитировать оригинальный email» в справочнике CaseNotificationRule позволяет '
    'включить цитату исходного письма (ParentActivity), но это корневое письмо, не последний ответ.'
)

doc.add_paragraph()
add_h(doc, '1.2 Текст ответа инженера в уведомлении заказчику', 2)
add_body(doc,
    'Текущее состояние: текст ответа инженера НЕ включается в email заказчику. '
    'Шаблон «Сообщение о получении нового ответа по обращению» содержит только номер/тему/статус.\n\n'
    'Возможное улучшение: invokable-макрос [#@Invoke.UsrLastEngineerReply#] в шаблоне. '
    'Запланировано как задача 2.6.'
)

doc.add_paragraph()
add_h(doc, '1.3 Имя отправителя', 2)
add_body(doc,
    'Текущее состояние: все письма уходят от servicedesk@cti.ru без подписи инженера.\n\n'
    'Возможное улучшение: добавить макрос [#Owner.Name#] в подвал шаблонов. '
    'Запланировано как задача 2.7.'
)

# ===================== РАЗДЕЛ 2 =====================
doc.add_page_break()
add_h(doc, '2. Уведомления в сторону Сотрудников (инженеры, сервис-менеджеры, 1-я линия)', 1)

add_h(doc, '2.1 Текущие механизмы (три разных, без единой конфигурации)', 2)
doc.add_paragraph()

t2 = doc.add_table(rows=1, cols=6)
t2.style = 'Table Grid'
add_header_row(t2, ['Событие', 'Получатель', 'Шаблон письма', 'Содержание', 'Механизм', 'Тип'])

emp_rows = [
    ('Назначен ответственным',
     'Назначенный сотрудник',
     'Назначение ответственного в обращении',
     '«Вас назначили ответственным по обращению №...»',
     'UsrSendEmailToSROwnerCustom1',
     'BPMN-процесс (CTI)'),
    ('Создано обращение с ответственным',
     'Ответственный сотрудник',
     'Назначение ответственного в обращении',
     'То же',
     'UsrSendEmailToSROwnerCustom1\n(второй StartSignal)',
     'BPMN-процесс (CTI)'),
    ('Назначено на группу\n(Owner=NULL, Group!=NULL)',
     'Все сотрудники группы',
     'На группу назначено обращение',
     '«На группу назначено обращение №...»',
     'RunSendEmailToCaseGroupV2',
     'BPMN-процесс (CaseService)'),
    ('Клиент написал ответ\n\u2192 «Получен ответ»',
     'Ответственный\n+ CC: роль «1-я линия»',
     'Добавление нового email\nпо обращению',
     '«Получен новый email по обращению №...»\n\u26a0 Текст письма клиента НЕ включён.',
     'UsrProcess_0c71a12CTI5',
     'BPMN-процесс (CTI)'),
    ('Сообщение с портала',
     'Ответственный',
     '(по шаблону)',
     '\u2014',
     'CasePortalMessageHistoryNotification\nProcessML',
     'BPMN (Portal)'),
]
for i, r in enumerate(emp_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t2, r, bg)
set_col_widths(t2, [Cm(3.5), Cm(3.0), Cm(5.0), Cm(5.0), Cm(5.5), Cm(3.5)])

doc.add_paragraph()
add_h(doc, '2.2 Критические GAP-ы в текущем механизме', 2)
doc.add_paragraph()

t_gap = doc.add_table(rows=1, cols=3)
t_gap.style = 'Table Grid'
add_header_row(t_gap, ['GAP', 'Описание', 'Влияние'])
gap_rows = [
    ('Нет текста письма клиента',
     'Инженер получает «Получен новый email...» без текста. Надо открывать систему.',
     '3-5 мин на каждое уведомление (задача 2.3)'),
    ('Фильтр по статусам',
     'UsrProcess_0c71a12CTI5 срабатывает ТОЛЬКО при смене статуса на «Получен ответ». '
     'Если обращение в статусе «Новое» или «В работе» — email инженеру не уходит (только push в колокольчик).',
     'Ответ клиента может быть пропущен (задача 2.3/2.4)'),
    ('Нет единого механизма',
     'Три разных BPMN на три события. Нет справочника правил (как CaseNotificationRule для клиентов).',
     'Сложно добавлять новые события, нет гибкой настройки'),
    ('Нет уведомлений о смене статуса',
     'Инженер не получает email при эскалации, переоткрытии, отмене и т.д.',
     'Задача 2.4'),
    ('Нет SLA-предупреждений',
     'Только визуальные индикаторы на странице.',
     'Задача 2.1'),
    ('Нет напоминаний о зависании',
     'Обращение может «зависнуть» в «Получен ответ».',
     'Задача 2.2'),
]
for i, r in enumerate(gap_rows):
    bg = 'FCE4D6' if i % 2 == 0 else 'FFF2CC'
    add_data_row(t_gap, r, bg)
set_col_widths(t_gap, [Cm(4.5), Cm(14), Cm(7)])

doc.add_paragraph()
add_h(doc, '2.3 UsrSendEmailToSROwnerCustom1 — диаграмма (по скриншоту 15.04.2026)', 2)
add_mono(doc,
    'StartSignal1: «Обращение назначено на ответственного» (смена Owner)\n'
    'StartSignal2: «Создано обращение с ответственным» (создание Case с Owner)\n'
    '  |\n'
    '  v\n'
    '[Задание-сценарий] -> [ExclusiveGateway]\n'
    '                          |\n'
    '        +-----------------+------------------+\n'
    '        | «Ответственный = Изменил»          | (другое условие -> ?)\n'
    '        v                                     |\n'
    '  [Читать данные обращения]                   |\n'
    '        |                                     |\n'
    '  [Читать данные ответственного]              |\n'
    '        |                                     |\n'
    '        +-- (нет email) -> [Terminate]        |\n'
    '        |                                     |\n'
    '        | «E-mail ответственного указан»      |\n'
    '        v                                     |\n'
    '  [Читать данные схемы Обращение]             |\n'
    '        |                                     |\n'
    '  [Обработать шаблон письма с макросами]      |\n'
    '        |                                     |\n'
    '  [ExclusiveGateway: «Происхождение по email?»]\n'
    '        |                                     |\n'
    '  +-----+----------------------+              |\n'
    '  | Не по email                | По email     |\n'
    '  v                            v              |\n'
    '  [Добавить данные email]  [Первая активность]|\n'
    '        |                  [Изменить тему]    |\n'
    '        |                       |             |\n'
    '        +----------+------------+             |\n'
    '                   v                          |\n'
    '             [Отправить письмо] <-------------+'
)
add_note(doc,
    'Ветка «Происхождение по email» — если обращение создано из email, процесс ищет первую Activity '
    '(корневое письмо) и меняет тему для email-threading (In-Reply-To). '
    'Групповое назначение здесь не обрабатывается — для группы работает RunSendEmailToCaseGroupV2.'
)

# ===================== РАЗДЕЛ 3 =====================
doc.add_page_break()
add_h(doc, '3. Push-уведомления внутри BPMSoft (не email)', 1)

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
add_header_row(t3, ['Событие', 'Получатель', 'Механизм (системное имя)', 'Тип механизма'])
push_rows = [
    ('Клиент написал по обращению\n(ЛЮБОЙ статус — без исключений)',
     'Ответственный инженер',
     'ReopenCaseAndNotifyAssignee \u2192 NotifyOwner()\n\u2192 запись Reminding в БД',
     'C#-класс (CaseService)'),
    ('Переоткрытие обращения (групповое)',
     'Группа ответственных',
     'Единое окно (штатный платформенный механизм)',
     'Платформа'),
]
for i, r in enumerate(push_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t3, r, bg)
set_col_widths(t3, [Cm(5.5), Cm(3.5), Cm(8), Cm(4)])

# ===================== РАЗДЕЛ 4 =====================
doc.add_paragraph()
add_h(doc, '4. Полная цепочка при получении email от клиента', 1)

steps = [
    ('1', 'BPMSoft получает входящее письмо',
     'Создаётся Activity (Type=Email, MessageType=Incoming). Письмо привязывается к обращению по номеру в теме.'),
    ('2', 'Сохранение CC-адресов',
     'UsrActivityCcEventListener (CTI) фиксирует CC-адреса из входящего письма в Case.UsrCcEmails.'),
    ('3', 'Запуск оркестратора',
     'Сигнал на новую Activity по обращению \u2192 BPMN RunSendNotificationCaseOwnerProcess (CaseService).'),
    ('4', 'Выбор пути (feature-toggle)',
     'toggle RunReopenCaseAndNotifyAssigneeClass=1 (на проде) \u2192 C#-класс ReopenCaseAndNotifyAssignee.\n'
     'При toggle=0 — шёл бы BPMN UsrSendNotificationToCaseOwnerCustom1 (на проде не используется).'),
    ('5', 'Переоткрытие (смена статуса)',
     'Предикат ReopeningCondition: если статус IsResolved или IsPaused \u2192 меняет на \u00abПолучен ответ\u00bb.\n'
     'Если \u00abНовое\u00bb, \u00abВ работе\u00bb, уже \u00abПолучен ответ\u00bb или финальный — статус НЕ меняется.'),
    ('6', 'Push-уведомление сотруднику (всегда)',
     'NotifyOwner() \u2192 запись Reminding \u2192 колокольчик BPMSoft. Срабатывает ВСЕГДА при наличии ответственного.'),
    ('7', 'Email сотруднику (только если статус изменился)',
     'Смена статуса на \u00abПолучен ответ\u00bb \u2192 StartSignal \u2192 UsrProcess_0c71a12CTI5 (CTI).\n'
     'Читает данные обращения и email ответственного \u2192 отправляет по шаблону.\n'
     '\u26a0 GAP: Текст клиентского письма в email НЕ включён.\n'
     '\u26a0 GAP: Если статус не изменился (обращение в \u00abНовое\u00bb/\u00abВ работе\u00bb) — email НЕ уходит.'),
]
for num, title_s, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run('Шаг {}. {}\n'.format(num, title_s))
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(9)
doc.add_paragraph()

# ===================== РАЗДЕЛ 5 =====================
add_h(doc, '5. Логика статуса \u00abПолучен ответ\u00bb \u2014 откуда берётся перечень статусов', 1)

add_h(doc, '5.1 Где задаётся в системе', 2)
add_body(doc,
    'Путь: Администрирование \u2192 Справочники \u2192 Состояния обращений (таблица CaseStatus)\n\n'
    'Каждое состояние имеет флаги:\n'
    '  \u2022 \u00abНа паузе\u00bb (IsPaused) \u2014 обращение ожидает чего-либо\n'
    '  \u2022 \u00abЯвляется решением\u00bb (IsResolved) \u2014 обращение считается решённым\n'
    '  \u2022 \u00abФинальный\u00bb (IsFinal) \u2014 обращение завершено (Закрыто/Отменено)\n'
    '  \u2022 \u00abПереоткрыто\u00bb (IsReopen) \u2014 статус \u00abПолучен ответ\u00bb\n\n'
    'Чтобы добавить новый статус в перечень \u2192 поставить ему флаг \u00abНа паузе\u00bb или '
    '\u00abЯвляется решением\u00bb в справочнике. Изменения в коде не требуются.'
)
doc.add_paragraph()

add_h(doc, '5.2 Предикат в C#-коде (ReopenCaseAndNotifyAssignee.cs, CaseService)', 2)
add_body(doc,
    'Переход в \u00abПолучен ответ\u00bb происходит, если выполнены ВСЕ три условия:\n'
    '  1. Статус НЕ является уже \u00abПолучен ответ\u00bb (!IsReopenStatus)\n'
    '  2. Статус НЕ финальный (!IsFinalStatus) \u2014 не \u00abЗакрыто\u00bb и не \u00abОтменено\u00bb\n'
    '  3. Статус является \u00abРешённым\u00bb ИЛИ \u00abНа паузе\u00bb (IsResolved OR IsPaused)\n\n'
    'Значения флагов читаются из справочника CaseStatus через CaseStatusStore \u2014 '
    'список не зашит в коде, берётся из БД.'
)
doc.add_paragraph()

add_h(doc, '5.3 Перечень статусов (с прода, апрель 2026)', 2)

t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
add_header_row(t4, ['Статус обращения', 'На паузе (IsPaused)', 'Решено (IsResolved)', 'Переход в \u00abПолучен ответ\u00bb?'])
status_rows = [
    ('Новое', '\u2014', '\u2014', 'Нет'),
    ('В работе', '\u2014', '\u2014', 'Нет'),
    ('Ожидает ответа', 'Да', '\u2014', 'Да'),
    ('Отложен', 'Да', '\u2014', 'Да'),
    ('Передано сервис-менеджеру', 'Да', '\u2014', 'Да'),
    ('Направлено в группу специализации', 'Да', '\u2014', 'Да'),
    ('Находится в работе у вендора', 'Да', '\u2014', 'Да'),
    ('Передано в отдел логистики', 'Да', '\u2014', 'Да'),
    ('Ожидаем поставщика', 'Да', '\u2014', 'Да'),
    ('Ожидается применение релиза', 'Да', '\u2014', 'Да'),
    ('Ожидает повторение проблемы', 'Да', '\u2014', 'Да'),
    ('Идёт доработка', 'Да', '\u2014', 'Да'),
    ('Идёт тестирование', 'Да', '\u2014', 'Да'),
    ('Передано в отдел разработки', 'Да', '\u2014', 'Да'),
    ('Передано в SD', 'Да', '\u2014', 'Да'),
    ('Приостановлен', 'Да', '\u2014', 'Да'),
    ('Решено', '\u2014', 'Да', 'Да'),
    ('Работы выполнены', '\u2014', 'Да', 'Да'),
    ('Получен ответ', '\u2014', '\u2014 (IsReopen=Да)', 'Нет (уже переоткрыт)'),
    ('Закрыто', '\u2014', '\u2014 (IsFinal=Да)', 'Нет (финальный)'),
    ('Отменено', '\u2014', '\u2014 (IsFinal=Да)', 'Нет (финальный)'),
]
for r in status_rows:
    v = r[3]
    bg = 'E2EFDA' if v == 'Да' else ('FCE4D6' if v.startswith('Нет') else 'FFF2CC')
    add_data_row(t4, r, bg)
set_col_widths(t4, [Cm(7), Cm(4), Cm(4), Cm(6)])

# ===================== РАЗДЕЛ 6 =====================
doc.add_page_break()
add_h(doc, '6. Полный список шаблонов уведомлений (по документации BPMSoft v1.9)', 1)
add_body(doc,
    'Шаблоны хранятся в справочнике \u00abШаблоны сообщений\u00bb (EmailTemplate). '
    'Для клиентских событий существуют пары шаблонов: один для Пользователя услуги, '
    'второй для Инициатора (суффикс \u00ab(для инициатора)\u00bb). '
    'Шаблоны редактируются в Дизайнере системы \u2192 Справочники \u2192 Шаблоны сообщений.'
)
doc.add_paragraph()

add_h(doc, 'Шаблоны для клиентов', 2)
t5a = doc.add_table(rows=1, cols=3)
t5a.style = 'Table Grid'
add_header_row(t5a, ['Шаблон', 'Вариант для инициатора', 'Назначение'])
client_tpl = [
    ('Создание нового обращения из чата', '\u2014', 'Уведомление о создании обращения через чат'),
    ('Сообщение о взятии обращения в работу', 'Сообщение о взятии обращения в работу (для инициатора)', 'Обращение взято в работу'),
    ('Сообщение о получении нового ответа по обращению', '...для инициатора', 'Инженер написал ответ по обращению'),
    ('Сообщение о разрешении обращения', '...для инициатора', 'Обращение решено'),
    ('Сообщение о разрешении обращения \u2014 только с решением', '...для инициатора', 'Обращение решено (упрощённый вариант)'),
    ('Подтверждение закрытия обращения', '...для инициатора', 'Запрос подтверждения закрытия'),
    ('Сообщение о закрытии обращения', '...для инициатора', 'Обращение закрыто'),
    ('Сообщение об отмене обращения', '...для инициатора', 'Обращение отменено'),
    ('Запрос оценки по обращению', '...для инициатора', 'Запрос оценки качества'),
]
for i, r in enumerate(client_tpl):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t5a, r, bg)
set_col_widths(t5a, [Cm(7), Cm(9), Cm(9.5)])

doc.add_paragraph()
add_h(doc, 'Шаблоны для сотрудников', 2)
t5b = doc.add_table(rows=1, cols=2)
t5b.style = 'Table Grid'
add_header_row(t5b, ['Шаблон', 'Назначение'])
emp_tpl = [
    ('Назначение ответственного в обращении', 'Внутреннее: ответственный назначен'),
    ('На группу назначено обращение', 'Внутреннее: обращение назначено на группу'),
    ('Добавление нового email по обращению', 'Внутреннее: получен новый email от клиента'),
]
for i, r in enumerate(emp_tpl):
    bg = 'E2EFDA' if i % 2 == 0 else None
    add_data_row(t5b, r, bg)
set_col_widths(t5b, [Cm(9), Cm(16.5)])

doc.add_paragraph()
add_h(doc, 'Служебные шаблоны', 2)
t5c = doc.add_table(rows=1, cols=2)
t5c.style = 'Table Grid'
add_header_row(t5c, ['Шаблон', 'Назначение'])
svc_tpl = [
    ('Пустой шаблон по обращению', 'Нестандартные уведомления (заготовка)'),
    ('Шаблон приглашения SSP', 'Приглашение клиента на портал'),
    ('Шаблон регистрации пользователя портала', 'Ссылка активации для портального пользователя'),
]
for i, r in enumerate(svc_tpl):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t5c, r, bg)
set_col_widths(t5c, [Cm(9), Cm(16.5)])

# ===================== РАЗДЕЛ 7 =====================
doc.add_page_break()
add_h(doc, '7. Полная инвентаризация процессов (с прода 15.04.2026)', 1)
add_body(doc, 'Источник: выгрузка представления vwprocesslib (все активные и неактивные процессы).')
doc.add_paragraph()

add_h(doc, '7.1 Уведомления клиенту', 2)
t7a = doc.add_table(rows=1, cols=5)
t7a.style = 'Table Grid'
add_header_row(t7a, ['Системное имя', 'Тип', 'Пакет', 'Статус на проде', 'Описание'])
client_proc = [
    ('SendEmailToCaseStatusChanged', 'C#-класс', 'CaseService', 'Активен (toggle=1)',
     'Единый механизм email клиенту при смене статуса. Читает CaseNotificationRule (статус+категория -> шаблон). Подставляет макросы.'),
    ('SendEmailToCaseStatusChangedProcess', 'BPMN', 'CaseService', 'Замещён C#-классом (toggle=1)',
     'BPMN-версия того же. Оркестратор: читает CaseNotificationRule, вызывает подпроцесс SendEmailToCaseContactPersonsProcess'),
    ('SendEmailToCaseContactPersonsProcess', 'BPMN', 'CaseService', 'Активен (подпроцесс)',
     'Хаб отправки клиенту. Разветвляет на 4 подпроцесса: контакт/инициатор x legacy/ML'),
    ('SendEmailToCaseContactProcessMultiLanguage', 'BPMN', 'CaseService', 'Активен (подпроцесс)',
     'Подпроцесс хаба: отправка контакту через EmailWithMacrosManager.SendEmail() (мультиязычный)'),
    ('SendEmailToCaseContactProcess', 'BPMN', 'Case', 'Legacy (подпроцесс)',
     'Подпроцесс хаба: отправка контакту (без мультиязычности)'),
    ('UsrProcess_send_reg_mail', 'BPMN', 'CTI', 'Активен',
     'Email заказчику о регистрации обращения (не покрывается CaseNotificationRule, т.к. создание != смена статуса)'),
    ('CasePortalMessageHistoryNotificationProcessMultiLanguage', 'BPMN', 'Portal', 'Активен',
     'Email клиенту + email сотруднику при сообщении с портала'),
    ('CasePortalMessageHistoryNotificationProcess', 'BPMN', 'Portal', 'Legacy-версия',
     'То же без мультиязычности'),
]
for i, r in enumerate(client_proc):
    bg = 'EBF3FB' if i % 2 == 0 else None
    if 'Замещён' in r[3] or 'Legacy' in r[3]:
        bg = 'FFF2CC'
    add_data_row(t7a, r, bg)
set_col_widths(t7a, [Cm(5.5), Cm(2), Cm(2.5), Cm(3), Cm(12.5)])

doc.add_paragraph()
add_h(doc, '7.2 Уведомления сотрудникам', 2)
t7b = doc.add_table(rows=1, cols=5)
t7b.style = 'Table Grid'
add_header_row(t7b, ['Системное имя', 'Тип', 'Пакет', 'Статус на проде', 'Описание'])
emp_proc = [
    ('UsrSendEmailToSROwnerCustom1', 'BPMN', 'CTI', 'Активен',
     'Email ответственному при назначении. Замещает SendEmailToSROwner. Два StartSignal: назначение + создание с ответственным. Ветвление «Происхождение по email» для threading'),
    ('SendEmailToSROwner', 'BPMN', 'CaseService', 'Замещён (CTI)',
     'Родитель UsrSendEmailToSROwnerCustom1. Не выполняется — замещён потомком'),
    ('RunSendEmailToCaseGroupV2', 'BPMN', 'CaseService', 'Активен',
     'Email всей группе при назначении. StartSignal: Case создан/изменён, Owner=NULL + Group!=NULL. Собирает email всех пользователей группы через SysUserInRole'),
    ('SendEmailToCaseGroup', 'BPMN', 'CaseService', 'Legacy',
     'Родитель RunSendEmailToCaseGroupV2. Вызывается как fallback'),
    ('UsrProcess_0c71a12CTI5', 'BPMN', 'CTI', 'Активен',
     'Email ответственному при переходе в «Получен ответ». StartSignal на изменение Case.StatusId. Шаблон: «Добавление нового email по обращению». CC: роль «1-я линия». CreateActivity=false'),
    ('ReopenCaseAndNotifyAssignee', 'C#-класс', 'CaseService', 'Активен (toggle=1)',
     'Push (Reminding) ответственному при ЛЮБОМ входящем email. Также меняет статус на «Получен ответ» (если IsResolved/IsPaused)'),
    ('RunSendNotificationCaseOwnerProcess', 'BPMN', 'CaseService', 'Активен (оркестратор)',
     'Оркестратор: при входящем email выбирает путь (C# или BPMN) по feature-toggle'),
    ('UsrSendNotificationToCaseOwnerCustom1', 'BPMN', 'CTI', 'НЕ активен (toggle=1)',
     'Замещает SendNotificationToCaseOwner. При toggle=1 обходится C#-классом. Функция: push + смена статуса'),
    ('SendNotificationToCaseOwner', 'BPMN', 'CaseService', 'Замещён (CTI)',
     'Родитель UsrSendNotificationToCaseOwnerCustom1'),
]
for i, r in enumerate(emp_proc):
    if 'НЕ активен' in r[3]:
        bg = 'FCE4D6'
    elif 'Замещён' in r[3] or 'Legacy' in r[3]:
        bg = 'FFF2CC'
    elif i % 2 == 0:
        bg = 'EBF3FB'
    else:
        bg = None
    add_data_row(t7b, r, bg)
set_col_widths(t7b, [Cm(5.5), Cm(2), Cm(2.5), Cm(3), Cm(12.5)])

doc.add_paragraph()
add_h(doc, '7.3 Оценка удовлетворённости (email клиенту)', 2)
t7c = doc.add_table(rows=1, cols=4)
t7c.style = 'Table Grid'
add_header_row(t7c, ['Системное имя', 'Тип', 'Пакет', 'Описание'])
sat_proc = [
    ('RunAnalyzeCaseSatisfactionLevel', 'BPMN', 'CaseService', 'Запуск процесса оценки (оркестратор)'),
    ('AnalyzeCaseSatisfactionLevel', 'BPMN', 'CaseService', 'Оценка удовлетворённости — отправляет запрос клиенту'),
    ('ReevaluateCaseLevelRequestProcessV2', 'BPMN', 'CaseService', 'Повторный запрос оценки'),
    ('ClearSatisfactionLevelProcess', 'BPMN', 'CaseService', 'Очистка поля «Уровень удовлетворенности» при решении'),
]
for i, r in enumerate(sat_proc):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t7c, r, bg)
set_col_widths(t7c, [Cm(6), Cm(2), Cm(2.5), Cm(15)])

doc.add_paragraph()
add_h(doc, '7.4 Обработка обращений (не отправляют уведомления, но связаны)', 2)
t7d = doc.add_table(rows=1, cols=4)
t7d.style = 'Table Grid'
add_header_row(t7d, ['Системное имя', 'Тип', 'Пакет', 'Описание'])
proc_other = [
    ('IncidentRegistrationFromEmailProcess', 'BPMN', 'CaseService', 'Регистрация обращения по входящему письму (создание Case)'),
    ('ActualizeCalculationTermsAfterStatusChange', 'BPMN', 'CaseService', 'Пересчёт SLA-сроков после смены статуса'),
    ('CaseOverduesSettingProcess', 'BPMN', 'CaseService', 'Установка показателей просроченности (визуальные индикаторы)'),
    ('IncidentDiagnosticsAndResolvingV2', 'BPMN', 'CaseService', 'Диагностика и решение инцидентов (создаёт задачу)'),
    ('AutoResolutionIncidentAfterClosedProblem', 'BPMN', 'CaseService', 'Авторешение инцидентов при закрытии проблемы'),
    ('PredictCaseFieldValuesProcess', 'BPMN', 'CaseService', 'ML-прогнозирование полей обращения'),
    ('StartSimilarCaseSearch / SimilarCasesSearch', 'BPMN', 'CaseService', 'Поиск похожих обращений'),
    ('SearchForParent', 'BPMN', 'CaseService', 'Поиск аналогичных инцидентов'),
    ('CreateCaseFromChat', 'BPMN', 'CaseService', 'Создание обращения из чата'),
    ('ActualizeCasePortalUserAction', 'BPMN', 'CaseService', 'Актуализация действий на портале'),
    ('SendResolution', 'BPMN', 'CaseService', 'Создание статьи в БЗ после решения (неактивен)'),
]
for i, r in enumerate(proc_other):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t7d, r, bg)
set_col_widths(t7d, [Cm(6), Cm(2), Cm(2.5), Cm(15)])

doc.add_paragraph()
add_h(doc, '7.5 Кастомные CTI-процессы (не уведомления)', 2)
t7e = doc.add_table(rows=1, cols=4)
t7e.style = 'Table Grid'
add_header_row(t7e, ['Системное имя', 'Тип', 'Статус', 'Описание'])
cti_other = [
    ('UsrProcess_a5f980e', 'BPMN', 'Неактивен', 'Изменение даты регистрации обращения'),
    ('UsrBindUserContact', 'BPMN', 'Активен', 'Привязка пользователя чата к контакту'),
]
for i, r in enumerate(cti_other):
    bg = 'FFF2CC' if r[2] == 'Неактивен' else ('EBF3FB' if i % 2 == 0 else None)
    add_data_row(t7e, r, bg)
set_col_widths(t7e, [Cm(5), Cm(2), Cm(2.5), Cm(16)])

# ===================== РАЗДЕЛ 8 =====================
doc.add_page_break()
add_h(doc, '8. Иерархия вызовов (полная схема)', 1)
add_mono(doc,
    '=== ПРИ СМЕНЕ СТАТУСА ОБРАЩЕНИЯ ===\n\n'
    '  Case.Status изменился\n'
    '    |\n'
    '    +-- toggle SendEmailToCaseOnStatusChangeClass=1 (наш случай)\n'
    '    |    -> C#-класс SendEmailToCaseStatusChanged\n'
    '    |         -> CaseNotificationRule -> шаблон -> email КЛИЕНТУ\n'
    '    |\n'
    '    +-- toggle=0\n'
    '         -> BPMN SendEmailToCaseStatusChangedProcess\n'
    '              -> SubProcess: SendEmailToCaseContactPersonsProcess (хаб)\n'
    '                   +-- SendEmailToCaseContactProcessMultiLanguage (контакт ML)\n'
    '                   +-- SendEmailToCaseContactProcess (контакт legacy)\n'
    '                   +-- SendEmailToInitiatorMultiLanguage (инициатор ML)\n'
    '                   +-- SendEmailToInitiator (инициатор legacy)\n\n'
    '=== ПРИ НАЗНАЧЕНИИ ОТВЕТСТВЕННОГО ===\n\n'
    '  Case.Owner изменился ИЛИ Case создан с Owner\n'
    '    -> UsrSendEmailToSROwnerCustom1 (CTI, замещает SendEmailToSROwner)\n'
    '         -> email ОТВЕТСТВЕННОМУ\n\n'
    '=== ПРИ НАЗНАЧЕНИИ НА ГРУППУ ===\n\n'
    '  Case создан/изменён, Owner=NULL, Group!=NULL\n'
    '    -> RunSendEmailToCaseGroupV2 (CaseService)\n'
    '         -> email ВСЕЙ ГРУППЕ\n\n'
    '=== ПРИ ОТВЕТЕ КЛИЕНТА (EMAIL) ===\n\n'
    '  Входящий email -> Activity -> сигнал\n'
    '    -> RunSendNotificationCaseOwnerProcess (оркестратор)\n'
    '         |\n'
    '         +-- toggle RunReopenCaseAndNotifyAssigneeClass=1 (наш случай)\n'
    '              -> C#-класс ReopenCaseAndNotifyAssignee.Run()\n'
    '                   |\n'
    '                   +-- ReopeningCondition: IsResolved || IsPaused\n'
    '                   |    -> StatusId = «Получен ответ» (если условие выполнено)\n'
    '                   |    -> статус НЕ меняется (если «Новое», «В работе» и т.д.)\n'
    '                   |\n'
    '                   +-- NotifyOwner() -> push (Reminding) — ВСЕГДА\n'
    '                   |\n'
    '                   +-- [Побочный эффект: смена статуса -> StartSignal]\n'
    '                        -> UsrProcess_0c71a12CTI5 (CTI)\n'
    '                             -> email ОТВЕТСТВЕННОМУ + CC: роль «1-я линия»\n'
    '                             -> GAP: текст клиентского письма НЕ включён\n'
    '                             -> GAP: срабатывает ТОЛЬКО если статус изменился\n\n'
    '=== ПРИ ОТВЕТЕ КЛИЕНТА (ПОРТАЛ) ===\n\n'
    '  PortalMessage создан\n'
    '    -> CasePortalMessageHistoryNotificationProcessMultiLanguage\n'
    '         +-- SubProcess: SendEmailToCaseContactPersonsProcess -> email КЛИЕНТУ\n'
    '         +-- SubProcess: CasePortalMessageHistoryNotificationProcess -> email СОТРУДНИКУ\n\n'
    '=== ПРИ РЕГИСТРАЦИИ ===\n\n'
    '  Case создан\n'
    '    -> UsrProcess_send_reg_mail (CTI) -> email КЛИЕНТУ\n\n'
    '=== ОЦЕНКА УДОВЛЕТВОРЁННОСТИ ===\n\n'
    '  Case -> «Решено»\n'
    '    -> RunAnalyzeCaseSatisfactionLevel -> AnalyzeCaseSatisfactionLevel\n'
    '         -> email КЛИЕНТУ (запрос оценки)'
)

# ===================== РАЗДЕЛ 9 =====================
doc.add_page_break()
add_h(doc, '9. Feature-toggles и системные настройки', 1)

add_h(doc, '9.1 Feature-toggles (с прода, апрель 2026)', 2)
t9a = doc.add_table(rows=1, cols=3)
t9a.style = 'Table Grid'
add_header_row(t9a, ['Code', 'State', 'Комментарий'])
ft_rows = [
    ('EmailMessageMultiLanguageV2', '1 (ВКЛ)', 'Конвейер EmailWithMacrosManager активен. Мультиязычные шаблоны'),
    ('UseAsyncEmailSender', '1 (ВКЛ)', 'Асинхронная отправка через AsyncEmailSender.SendAsync()'),
    ('DelayedNotification', '1 (ВКЛ)', 'Отложенные уведомления'),
    ('SendEmailToCaseOnStatusChangeClass', '1 (ВКЛ)', 'Email клиенту при смене статуса — C#-класс вместо BPMN'),
    ('RunReopenCaseAndNotifyAssigneeClass', '1 (ВКЛ)', 'Переоткрытие — C#-класс вместо BPMN'),
]
for i, r in enumerate(ft_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t9a, r, bg)
set_col_widths(t9a, [Cm(7), Cm(3), Cm(15.5)])

doc.add_paragraph()
add_h(doc, '9.2 Системные настройки', 2)
t9b = doc.add_table(rows=1, cols=3)
t9b.style = 'Table Grid'
add_header_row(t9b, ['Code', 'Значение', 'Комментарий'])
ss_rows = [
    ('SiteUrl', 'bpm.cti.ru', 'Базовый URL для ссылок в email'),
    ('SupportServiceEmail', 'servicedesk@cti.ru', 'Адрес отправителя'),
    ('ClearAssigneeOnCaseReopening', 'false', 'При переоткрытии ответственный не сбрасывается'),
]
for i, r in enumerate(ss_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t9b, r, bg)
set_col_widths(t9b, [Cm(7), Cm(4), Cm(14.5)])

# ===================== РАЗДЕЛ 10 =====================
doc.add_page_break()
add_h(doc, '10. Архитектурное решение: UsrEmployeeNotificationManager', 1)

add_h(doc, '10.1 Проблема', 2)
add_body(doc,
    'Для клиентов существует единый структурированный механизм: C#-класс SendEmailToCaseStatusChanged + '
    'справочник CaseNotificationRule (статус + категория -> шаблон). '
    'Для сотрудников такого механизма нет — есть три разрозненных BPMN-процесса, '
    'привязанных к конкретным событиям, без общей конфигурации.'
)

doc.add_paragraph()
add_h(doc, '10.2 Решение', 2)
add_body(doc,
    'Создать единый C#-класс UsrEmployeeNotificationManager + справочник правил '
    'UsrEmployeeNotificationRule — по аналогии с механизмом для клиентов.'
)

doc.add_paragraph()
add_h(doc, '10.3 Справочник UsrEmployeeNotificationRule', 2)
t10a = doc.add_table(rows=1, cols=3)
t10a.style = 'Table Grid'
add_header_row(t10a, ['Поле', 'Тип', 'Назначение'])
rule_rows = [
    ('Событие (UsrEvent)', 'Справочник',
     'Тип события: ОтветКлиента, НазначениеОтветственного, НазначениеГруппы, СменаСтатуса, SLA-порог, Зависание'),
    ('Статус обращения (StatusId)', 'Справочник (CaseStatus)',
     'Для каких статусов срабатывает (NULL = любой)'),
    ('Шаблон email (EmailTemplateId)', 'Справочник (EmailTemplate)',
     'Какой шаблон использовать'),
    ('Получатель (UsrRecipientType)', 'Справочник',
     'Ответственный / Группа / Руководитель группы / Наблюдатели'),
    ('Правило использования', 'Справочник',
     'Отправляется сразу / С задержкой / Не используется'),
    ('Задержка, мин', 'Integer', 'Пауза перед отправкой'),
    ('Активно', 'Boolean', 'Вкл/Выкл'),
]
for i, r in enumerate(rule_rows):
    bg = 'EBF3FB' if i % 2 == 0 else None
    add_data_row(t10a, r, bg)
set_col_widths(t10a, [Cm(5.5), Cm(4.5), Cm(15.5)])

doc.add_paragraph()
add_h(doc, '10.4 Что заменяет новый механизм', 2)
t10b = doc.add_table(rows=1, cols=2)
t10b.style = 'Table Grid'
add_header_row(t10b, ['Текущий механизм', 'Что станет'])
replace_rows = [
    ('UsrProcess_0c71a12CTI5 (email при «Получен ответ»)', 'Правило: событие=ОтветКлиента, статус=любой'),
    ('UsrSendEmailToSROwnerCustom1 (email при назначении)', 'Правило: событие=НазначениеОтветственного'),
    ('RunSendEmailToCaseGroupV2 (email группе)', 'Правило: событие=НазначениеГруппы'),
    ('(нет)', 'Правило: событие=СменаСтатуса (задача 2.4)'),
    ('(нет)', 'Правило: событие=SLA-порог (задача 2.1)'),
    ('(нет)', 'Правило: событие=Зависание (задача 2.2)'),
]
for i, r in enumerate(replace_rows):
    has_new = r[0] == '(нет)'
    bg = 'E2EFDA' if has_new else ('EBF3FB' if i % 2 == 0 else None)
    add_data_row(t10b, r, bg)
set_col_widths(t10b, [Cm(9), Cm(16.5)])

doc.add_paragraph()
add_h(doc, '10.5 Точки вызова', 2)
add_body(doc,
    '  \u2022 EventListener на Case (OnUpdated) — смена статуса, смена ответственного\n'
    '  \u2022 EventListener на Activity (OnInserted) — входящий email (заменяет косвенную связь через смену статуса)\n'
    '  \u2022 Фоновый процесс (BPMN-таймер / CronJob) — SLA-пороги, зависание'
)

# ===================== СОХРАНЕНИЕ =====================
out_path = r'c:\Users\echum\Documents\BPMsoft\projects\notifications-wave2\email_notifications_description.docx'
doc.save(out_path)
print('Saved: ' + out_path)
